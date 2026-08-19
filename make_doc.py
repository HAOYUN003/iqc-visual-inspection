# -*- coding: utf-8 -*-
"""
重写 IQC 智能来料检测系统.docx 为正式系统说明。
整合后半段草稿为正式章节，修复过时信息，统一标题层级，保留表格与图片。
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 默认字体
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(11)
from docx.oxml.ns import qn
style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

ACCENT = RGBColor(0x25, 0x63, 0xEB)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = ACCENT
        r.font.name = "微软雅黑"
    return p


def h2(text):
    return doc.add_heading(text, level=2)


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    return doc.add_paragraph(text, style="List Number")


# ============ 封面标题 ============
title = doc.add_heading("AI + 机器视觉 + IQC 来料质量管理系统", level=0)
for r in title.runs:
    r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
r = sub.add_run("—— 面向半导体设备厂的智能视觉检测与质量追溯系统说明")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_paragraph()

# ============ 一、系统概述 ============
h1("一、系统概述")
para("本系统面向半导体设备厂来料检验（IQC）场景，融合光学成像、机器视觉与深度学习技术，"
     "构建了集规格防错、表面缺陷检测、尺寸判定、质量追溯于一体的智能检测平台。"
     "系统采用“AI 初筛 + 工程师决策”模式，在降低人工工作量的同时保证判定一致性，"
     "并逐步沉淀企业质量数据资产。")

h2("1.1 检测能力")
table = doc.add_table(rows=6, cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["功能", "具体内容", "判定方式"]
rows = [
    ["标准件规格防错", "读卡尺照片识别规格 M3~M8，与料单期望比对", "尺寸反推 / CNN 分类，防拿错混料"],
    ["加工件表面缺陷", "划伤、毛刺、碰伤、裂纹、锈蚀、麻点", "视觉大模型多角度判定，多张照片聚合"],
    ["螺纹状态检测", "侧视图识别螺纹正常 / 缺牙 / 烂牙", "本地 CNN 分类模型"],
    ["图纸清单逐项校验", "按检验标准逐项判定（部位 + 图纸阈值）", "表面类视觉判定；尺寸类按公差量化判定"],
    ["检验标准管理", "主管端录入检验部位 + 图纸阈值，按料号自动匹配", "结构化数据，可追溯、可演进"],
]
for j, htext in enumerate(headers):
    table.rows[0].cells[j].text = htext
for i, row in enumerate(rows, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

doc.add_paragraph()
h2("1.2 追溯与管理能力")
bullet("质检数据库（SQLite）：批次 / 供应商 / 料号 / 检验员 / 照片 / 结果全程留痕")
bullet("主管端报表：供应商质量排行、缺陷分布、物料良率趋势、复核一致性、抽检覆盖率")
bullet("上机不良反查：输入批次号一键反查该批次全部抽检记录与照片，快速定位漏检")
bullet("人工复核闭环：AI 初筛 → 工程师复核 → 复核结果回灌模型持续优化")
bullet("检测单打印：合格 / 不合格单据一键生成并打印，免手写")

# ============ 二、技术路线 ============
h1("二、技术路线")
para("检测流程：")
numbered("图像采集：手机 / 摄像头 / 卡尺拍照，含拍照规范指引与图像质量自动检查")
numbered("视觉分析：规格识别（尺寸反推 + CNN）、缺陷判定（视觉大模型 / 本地 YOLO）、螺纹分类")
numbered("判定逻辑：规格反推 + 尺寸公差判定 + 缺陷判定 → OK / NG / UNSURE")
numbered("数据闭环：判定入库 → 人工复核 → 复核数据导出训练集 → 本地模型重训")

para("")
para("技术组成：", bold=True)
bullet("视觉理解引擎：Qwen-VL-plus（阿里云，读卡尺 LCD / 判缺陷）")
bullet("本地模型（RTX 4060 免费推理）：ResNet18 规格 CNN、YOLO 缺陷模型、螺纹分类 CNN、CV 尺寸测量")
bullet("界面：Streamlit 双端（检验员端 / 主管端）")
bullet("数据存储：SQLite 质量数据库")

# ============ 三、核心优势 ============
h1("三、核心优势")
bullet("诚实性原则：视觉无法验证的项（硬度、粗糙度、精密尺寸）明确标注“需仪器检测”，不瞎猜")
bullet("AI 初筛 + 人工复核：不是替代检验员，而是降低工作量、提高判定一致性")
bullet("检验标准结构化：检验部位、图纸阈值存为数据而非提示词文字，可追溯、可演进")
bullet("聚焦重点物料：物料白名单机制，对稳定供应的重要来料优先建立数据资产")
bullet("可解释：每个判定带原因（偏差多少 mm、缺陷几处 vs 上限），非黑箱")
bullet("成本可控：规格 / 尺寸 / 螺纹走本地免费模型，仅表面缺陷调用视觉大模型（约 0.02 元 / 件）")

# ============ 五、系统使用流程 ============
h1("四、系统使用流程")

h2("4.1 检验员端：现场检验")
para("检验员端负责来料现场检验，核心流程如下：", bold=False)

para("（1）登记来料信息", bold=True)
para("填写批次号、料号、名称、期望规格、供应商、批次数量、检验员，然后点击「登记批次」。"
     "批次号是系统数据主键，后续照片、AI 结果、检验记录均绑定到该批次。"
     "支持扫码枪快速录入批次条码。"
     "登记后系统按料号自动匹配该物料的检验标准与图纸。")

para("（2）查阅图纸", bold=True)
para("登记料号后，系统自动关联该料号的图纸文件（PDF / 图片），检验员可在线预览或下载，"
     "对照图纸核对规格与技术要求。若无图纸，可提醒主管端上传。")

para("（3）图像采集", bold=True)
para("系统提供三种检测模式：")
bullet("标准件（卡尺读数定规格）：主要判断尺寸是否符合标准")
bullet("加工件（多角度表面缺陷）：偏向机器视觉外观检测（划痕、毛刺、缺口、崩角、污染、锈蚀等）")
bullet("螺纹（侧视图）：识别螺纹正常 / 缺牙 / 烂牙")
para("拍照前系统提示规范（零件水平、占画面 60%~80%、卡尺同框、光照均匀），"
     "并对上传图片自动质检（过暗 / 模糊 / 零件过小会提醒重拍）。")

para("（4）检测与结果判定", bold=True)
para("点击「开始检测」后，系统输出判定结果，示例：")
res_table = doc.add_table(rows=7, cols=2)
res_table.style = "Table Grid"
res_rows = [
    ("项目", "结果"),
    ("AI 判定", "OK / NG / 需复检"),
    ("规格识别", "M4"),
    ("尺寸", "4.02 mm（名义 4.00 ± 0.05 mm）"),
    ("螺纹状态", "正常 / 缺牙 / 烂牙"),
    ("表面缺陷", "无 / 划痕等"),
    ("判定原因", "偏差值、缺陷数量 vs 上限等可解释说明"),
]
for i, (a, b) in enumerate(res_rows):
    res_table.rows[i].cells[0].text = a
    res_table.rows[i].cells[1].text = b

para("")
para("（5）入库、复核与打印", bold=True)
para("检测结果经工程师复核后入库（AI 初筛 + 工程师决策）；"
     "合格 / 不合格可一键生成检测单并打印，免手写。检测记录关联图纸，全程可追溯。")

h2("4.2 主管端：质量统计与追溯")
para("主管端提供质量总览与追溯能力：")
bullet("质量总览：检验记录总数、OK/NG 分布、缺陷类型分布、物料良率趋势、复核一致性")
bullet("批次 / 供应商报表：批次明细、抽检覆盖率、供应商质量排名（可按 OK 率排序）")
bullet("追溯查询：按批次号 / 料号 / 供应商 / 日期检索，支持导出")
bullet("上机反查：输入批次号反查该批次全部抽检记录与照片，定位漏检")
bullet("物料管理：重点物料白名单配置（是否启用 AI、检测项开关）")
bullet("图纸管理：上传图纸文件（PDF/图片），按图号/料号关联，供检验员端查阅")
bullet("检验标准：录入检验部位 + 图纸阈值（含国标未注公差自动查表），检验员端按料号自动匹配加载")
bullet("数据积累：显示规格 / 读数配对 / 缺陷样本积累进度，达标后可重训本地模型")

para("")
para("供应商质量排名示例：")
sup_table = doc.add_table(rows=4, cols=4)
sup_table.style = "Table Grid"
sup_rows = [
    ("供应商", "来料批次", "OK 率", "NG 率"),
    ("A 供应商", "38", "99.2%", "0.8%"),
    ("B 供应商", "27", "96.4%", "3.6%"),
    ("C 供应商", "42", "99.8%", "0.2%"),
]
for i, row in enumerate(sup_rows):
    for j, val in enumerate(row):
        sup_table.rows[i].cells[j].text = val

para("")
para("质量追溯的意义：出现质量事故时，可快速反查——哪个供应商、哪个批次、什么料号、"
     "何时入库、谁检验、AI 判定结果、有无缺陷图片，实现工业质量系统要求的 Traceability。")

# ============ 六、当前问题与解决方案 ============
h1("五、当前问题与解决方案")

h2("5.1 来料种类多、数量不稳定，数据集不足")
para("现状：来料种类繁多且数量不稳定，系统缺乏足够数据集对所有来料进行误差检验。")
para("解决方案：聚焦供应商稳定发货、数量较大的来料，将其纳入系统白名单；"
     "通过持续积累这些来料的检验数据（人工复核回灌），使系统逐步学习，目标准确率达到 99% 以上。")

h2("5.2 人工抽检效率限制，存在漏检风险")
para("现状：主要依靠人工抽检，无法对来料逐一检验，可能出现抽检良品但袋中混有不良品、"
     "上机时才发现的情况。")
para("解决方案：系统记录每次抽检的批次、照片与判定，提供“上机不良反查”功能——"
     "上机发现问题后输入批次号，即可反查该批次全部抽检记录与照片，快速定位是否漏检；"
     "同时统计抽检覆盖率，为合理设定抽检比例提供依据。")

h2("5.3 检验环境不规范，尺寸计算易失真")
para("现状：任意角度、距离拍照导致比例压缩，尺寸计算可能失真。")
para("解决方案：系统内置拍照规范指引（零件水平摆放、占画面 60%~80%、卡尺同框、光照均匀），"
     "并对上传图片自动质检（过暗 / 模糊 / 零件占比过小会提醒重拍）。"
     "后续 Phase 1 引入工业相机 + 固定工位 + 光源，实现像素级精确定位与测量。")


h2("5.4 检验记录依赖手写，效率低")
para("现状：检验合格 / 不合格单据需手写，效率低且不易追溯。")
para("解决方案：系统判断合格 / 不合格后，一键生成可打印检测单（含批次、料号、规格、判定、"
     "缺陷明细、检验员、时间），浏览器直接打印，免手写。")

h2("5.5 视觉大模型读卡尺成本")
para("现状：视觉大模型读卡尺约 0.02 元 / 件，精度可靠但非零成本。")
para("解决方案：检验员复核时读数自动入库，累积“卡尺图 + 读数”配对数据；"
     "攒够一定数量（目标 300 对）后训练本地读数小模型，逐步替代视觉大模型，实现零成本。")

h2("5.6 视觉深度与测量精度")
para("现状：标准件有厚度，不可能完全贴在工作台上，零件在视觉上会抬离一定高度，"
     "导致相机画面比例变化；且目前无专业 CCD 相机，测量精度受限。")
para("光学原理与解决方案：")
numbered("澄清误区：当零件顶面平行于相机像平面时，透视投影是等比缩放，不是压缩变形——"
         "高度只改变整体缩放系数（物距），不产生形状畸变。真正的问题不是“高度导致压缩”，"
         "而是“高度不确定”导致比例不确定。")
numbered("治具固定高度：让零件顶面每次都落在同一已知高度 → 物距恒定 → 标定一次、比例固定，"
         "消除高度漂移引起的缩放误差。")
numbered("固定工位 + 垂直正拍：光轴垂直于工作台、零件居中，消除偏心畸变与斜视畸变。")
numbered("远心镜头（Phase 1 标配）：平行投影成像，放大倍率与物距无关，彻底免疫零件厚度变化——"
         "厚度件的标准测量方案。")
numbered("标定补偿：测量顶面实际高度，反算物距修正比例（软方案，配合治具使用）。")
numbered("CCD 相机 + 标定板：Phase 1 引入工业相机与标定板，亚像素测量，"
         "把判定公差从 ±0.4mm 收窄到 ±0.01mm 级。")
h2("5.7 工件上下两面尺寸不一致")
para("现状：部分工件上下两面尺寸不一致（例如一个面 15cm、另一面 15.2cm），"
     "单视角俯拍只能看到顶面，无法判断底面是否符合要求。")
para("解决方案：")
numbered("先读图定位要求：确认图纸公差标注的是顶面、底面还是某处截面——决定测量目标。")
numbered("多视角检测：正面 + 侧面 + 孔位多角度拍照（现有加工件模式已支持），"
         "让视觉模型判断各面状态。")
numbered("侧面相机 / 结构光 / 激光线扫（Phase 1+）：直接测量 3D 轮廓，覆盖上下尺寸差异。")
numbered("判定依据是图纸公差带而非“上下一致”：锥形件若图纸允许锥度，"
         "15cm 与 15.2cm 均可能合格——系统判断“是否在公差带内”，而非“上下是否一样”。")
h2("5.8 图纸上传与查阅")
para("现状：检验员端无法上传/查看图纸，图纸规格需依赖主管端预先录入检验标准。")
para("解决方案：新增图纸管理——主管端上传图纸文件（PDF/图片），按图号/料号关联；"
     "检验员端登记料号后自动匹配图纸，可在线预览或下载，检测记录关联图纸，实现可追溯。")

h1("六、后期改进路径")
numbered("固定工位 + 治具 + 光源（Phase 1，核心）：零件固定摆放 → 像素级 ROI 生效，"
         "低角度光源提升划痕对比度，减少视觉误判（详见《光学方案设计》）")
numbered("本地模型替代：读数 → 本地小模型；缺陷 → 真实数据重训 YOLO；规格 → ResNet18；"
         "视觉大模型仅保留疑难复核")
numbered("工业相机 + 标定：亚像素精密尺寸测量，把判定公差从 ±0.4mm 收窄到 ±0.01mm 级")
numbered("真实数据积累回灌：重点物料持续采集 → 微调本地模型 → 精度随数据增长")
numbered("对接公司现有系统：图纸型号从原系统自动带出，无需单独输入（待对接）")
numbered("与仓库联动：检验合格后自动放行通知、批次信息联动（待规划）")
numbered("对接 MES / ERP：检验结果自动推送、批次放行联动、供应链质量闭环")

# ============ 八、待确认事项 ============
h1("七、待确认事项")
para("以下事项需结合公司实际环境确定后推进：", italic=True)
bullet("对接现有来料检验系统：确认原系统的数据访问方式（网页 / 软件 / Excel、有无 API），"
       "实现图纸型号自动带出")
bullet("仓库联动方案：明确联动范围（库存核对 / 放行通知 / 扫码登记）")
bullet("检测单格式：对齐公司现有单据模板与栏位")
bullet("数据安全：图纸、检验记录等涉密资料的访问控制与存储安全")

# ============ 八、GitHub 仓库与登录指引 ============
h1("八、代码托管与 GitHub 登录指引")

h2("8.1 GitHub 仓库")
para("本项目代码开源托管在 GitHub，包含全部源码、README 使用说明与配置，便于版本管理与协作：")
bullet("仓库地址：https://github.com/HAOYUN003/iqc-visual-inspection")
bullet("内容：前端界面（app/）、检测引擎（src/）、配置、README")
bullet("安全说明：API 密钥（.mcp.json）、真实照片、模型权重、数据库均通过 .gitignore 排除，"
       "不随仓库公开。")

h2("8.2 首次登录 GitHub")
para("访问并克隆代码需先登录 GitHub 账号，步骤如下：", bold=False)
numbered("打开浏览器，访问 https://github.com ，点击右上角 Sign in 登录（无账号则点 Sign up 注册）")
numbered("进入仓库页面 https://github.com/HAOYUN003/iqc-visual-inspection ，"
         "点击绿色 Code 按钮，选择 HTTPS 地址复制")
numbered("在本地终端执行 git clone https://github.com/HAOYUN003/iqc-visual-inspection.git "
         "即可拉取代码到本地")

h2("8.3 提交代码 / 更新仓库（开发者）")
numbered("本地安装 GitHub CLI（gh）：Windows 可在 PowerShell 运行 winget install GitHub.cli")
numbered("登录：终端执行 gh auth login，按提示选择 GitHub.com → HTTPS → 网页授权")
numbered("提交本地改动：git add . → git commit -m \"说明\" → git push origin main")
para("提示：系统已配置 git 代理（127.0.0.1:7890），联网需开启代理，否则推送可能超时。",
     italic=True)

# ============ 图片（架构图，从原文档提取）============
import os
import zipfile
_img_file = "_doc_image1.png"
if not os.path.exists(_img_file):
    try:
        with zipfile.ZipFile("IQC 智能来料检测系统.docx") as zf:
            for n in zf.namelist():
                if n.startswith("word/media") and n.endswith(".png"):
                    with open(_img_file, "wb") as f:
                        f.write(zf.read(n))
                    break
    except Exception:
        pass
if os.path.exists(_img_file):
    doc.add_page_break()
    h1("附录：系统架构图")
    doc.add_picture(_img_file, width=Cm(15))

doc.save("IQC 智能来料检测系统（优化版）.docx")
print("文档已生成: IQC 智能来料检测系统（优化版）.docx")
print("段落数:", len(doc.paragraphs), "| 表格数:", len(doc.tables))
